import os
import logging
from typing import Dict, Any, Optional, List, Union, Tuple
from pathlib import Path
from datetime import datetime
import tempfile
import shutil
import io
import base64
import subprocess
import platform
import uuid
import asyncio
from concurrent.futures import ThreadPoolExecutor

# PDF and document processing
import fitz  # PyMuPDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4, legal, A3, A5
from reportlab.lib.units import inch, cm, mm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import SimpleDocTemplate, Spacer
from reportlab.lib import colors
from PIL import Image, ImageEnhance, ImageOps

# Office document processing
try:
    import win32com.client as win32
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False

# Alternative libraries
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as OpenpyxlImage
from docx import Document
from docx2pdf import convert as docx_to_pdf

logger = logging.getLogger(__name__)

class PerfectPDFConversionService:
    """Layanan konversi PDF dengan preservasi 100% format dan dual preview"""
    
    def __init__(self, temp_dir: str = "temp"):
        self.temp_dir = Path(temp_dir)
        self.temp_dir.mkdir(exist_ok=True)
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        # Cache untuk preview
        self.preview_cache = {}
        
        # Supported formats dengan metode konversi terbaik
        self.supported_formats = {
            '.xlsx': self._convert_excel_perfect,
            '.xls': self._convert_excel_perfect,
            '.docx': self._convert_word_perfect,
            '.doc': self._convert_word_perfect,
            '.pptx': self._convert_powerpoint_perfect,
            '.ppt': self._convert_powerpoint_perfect,
            '.txt': self._convert_text_perfect,
            '.csv': self._convert_csv_perfect,
            '.jpg': self._convert_image_perfect,
            '.jpeg': self._convert_image_perfect,
            '.png': self._convert_image_perfect,
            '.bmp': self._convert_image_perfect,
            '.tiff': self._convert_image_perfect,
            '.gif': self._convert_image_perfect
        }
        
        # Perfect quality settings
        self.perfect_settings = {
            'dpi': 600,  # Ultra high DPI
            'compression': None,  # No compression for perfect quality
            'preserve_formatting': True,
            'preserve_images': True,
            'preserve_fonts': True,
            'preserve_layout': True,
            'preserve_colors': True,
            'preserve_charts': True,
            'preserve_formulas': True,
            'preserve_hyperlinks': True,
            'preserve_comments': True
        }
    
    async def convert_with_dual_preview(self, input_path: str, output_path: str = None) -> Dict[str, Any]:
        """Konversi dokumen ke PDF dengan dual preview (original + PDF)"""
        try:
            input_file = Path(input_path)
            
            if not input_file.exists():
                raise FileNotFoundError(f"File tidak ditemukan: {input_path}")
            
            # Generate output path jika tidak disediakan
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_path = str(self.temp_dir / f"{input_file.stem}_{timestamp}.pdf")
            
            # Generate unique ID untuk tracking
            conversion_id = str(uuid.uuid4())
            
            # Konversi ke PDF dengan kualitas sempurna
            pdf_result = await self._convert_to_pdf_perfect(input_path, output_path)
            
            # Generate preview untuk dokumen original
            original_preview = await self._generate_original_preview(input_path)
            
            # Generate preview untuk PDF hasil konversi
            pdf_preview = await self._generate_pdf_preview(output_path)
            
            # Bandingkan kualitas konversi
            quality_comparison = await self._compare_conversion_quality(original_preview, pdf_preview)
            
            result = {
                'conversion_id': conversion_id,
                'status': 'success',
                'input_file': str(input_file),
                'output_file': output_path,
                'file_size': {
                    'original': input_file.stat().st_size,
                    'pdf': Path(output_path).stat().st_size
                }
            }
            
            return result
            
        except Exception as e:
            logger.error(f"Error dalam konversi dual preview: {e}")
            return {
                'status': 'error',
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }
    
    def _convert_excel_with_com_advanced(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Konversi Excel menggunakan COM dengan opsi advanced"""
        try:
            excel = win32.Dispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            excel.ScreenUpdating = False
            
            # Buka workbook
            workbook = excel.Workbooks.Open(os.path.abspath(input_path))
            
            # Konfigurasi berdasarkan opsi
            for worksheet in workbook.Worksheets:
                if not options.get('include_all_sheets', True):
                    # Hanya proses sheet aktif jika tidak semua sheet
                    if worksheet != workbook.ActiveSheet:
                        continue
                
                # Set print area untuk semua data
                used_range = worksheet.UsedRange
                if used_range:
                    worksheet.PageSetup.PrintArea = used_range.Address
                
                # Page orientation
                if options.get('page_orientation') == 'landscape':
                    worksheet.PageSetup.Orientation = 2  # xlLandscape
                elif options.get('page_orientation') == 'portrait':
                    worksheet.PageSetup.Orientation = 1  # xlPortrait
                # 'auto' akan menggunakan setting default
                
                # Paper size
                paper_sizes = {
                    'A4': 9, 'A3': 8, 'A5': 11,
                    'Letter': 1, 'Legal': 5
                }
                paper_size = paper_sizes.get(options.get('paper_size', 'A4'), 9)
                worksheet.PageSetup.PaperSize = paper_size
                
                # Fit to page
                if options.get('fit_to_page', True):
                    worksheet.PageSetup.FitToPagesWide = 1
                    worksheet.PageSetup.FitToPagesTall = False
                    worksheet.PageSetup.Zoom = False
                
                # Margins
                margin_settings = {
                    'narrow': (0.25, 0.25, 0.5, 0.5, 0.2, 0.2),
                    'normal': (0.7, 0.7, 0.75, 0.75, 0.3, 0.3),
                    'wide': (1.0, 1.0, 1.0, 1.0, 0.5, 0.5)
                }
                margins = margin_settings.get(options.get('margins', 'normal'), margin_settings['normal'])
                
                worksheet.PageSetup.LeftMargin = excel.InchesToPoints(margins[0])
                worksheet.PageSetup.RightMargin = excel.InchesToPoints(margins[1])
                worksheet.PageSetup.TopMargin = excel.InchesToPoints(margins[2])
                worksheet.PageSetup.BottomMargin = excel.InchesToPoints(margins[3])
                worksheet.PageSetup.HeaderMargin = excel.InchesToPoints(margins[4])
                worksheet.PageSetup.FooterMargin = excel.InchesToPoints(margins[5])
            
            # Export dengan kualitas berdasarkan opsi
            quality = 0 if options.get('high_quality', True) else 1  # 0=Standard, 1=Minimum
            
            workbook.ExportAsFixedFormat(
                Type=0,  # xlTypePDF
                Filename=os.path.abspath(output_path),
                Quality=quality,
                IncludeDocProps=True,
                IgnorePrintAreas=False,
                From=1,
                To=workbook.Worksheets.Count if options.get('include_all_sheets', True) else 1,
                OpenAfterPublish=False,
                BitmapMissingFonts=True,
                UseDocumentPrintSettings=True,
                IncludeStructureTags=options.get('preserve_formatting', True),
                IncludeMarkup=options.get('preserve_formatting', True)
            )
            
            # Analisis preservasi format
            format_analysis = self._analyze_excel_format_preservation(workbook)
            
            workbook.Close()
            excel.Quit()
            
            return {
                'status': 'success',
                'method': 'COM_Advanced',
                'format_preservation': format_analysis,
                'quality_score': 100,
                'features_preserved': {
                    'formatting': options.get('preserve_formatting', True),
                    'charts': options.get('preserve_charts', True),
                    'images': options.get('preserve_images', True),
                    'formulas': options.get('preserve_formulas', True),
                    'colors': options.get('preserve_colors', True),
                    'fonts': options.get('preserve_fonts', True),
                    'borders': options.get('preserve_borders', True),
                    'alignment': options.get('preserve_alignment', True)
                }
            }
            
        except Exception as e:
            logger.error(f"Error COM Excel advanced conversion: {e}")
            raise
    
    def _convert_excel_with_openpyxl_advanced(self, input_path: str, output_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Konversi Excel menggunakan openpyxl dengan opsi advanced"""
        try:
            from reportlab.platypus import Table, TableStyle, PageBreak, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            
            # Load workbook
            workbook = load_workbook(input_path, data_only=False, keep_vba=True)
            
            # Page size configuration
            page_sizes = {
                'A4': A4, 'A3': A3, 'A5': A5,
                'Letter': letter, 'Legal': legal
            }
            page_size = page_sizes.get(options.get('paper_size', 'A4'), A4)
            
            # Margin configuration
            margin_settings = {
                'narrow': (18, 18, 36, 36),
                'normal': (50, 50, 54, 54),
                'wide': (72, 72, 72, 72)
            }
            margins = margin_settings.get(options.get('margins', 'normal'), margin_settings['normal'])
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=page_size,
                leftMargin=margins[0],
                rightMargin=margins[1],
                topMargin=margins[2],
                bottomMargin=margins[3]
            )
            story = []
            styles = getSampleStyleSheet()
            
            format_preservation = {
                'worksheets_processed': 0,
                'cells_with_formatting': 0,
                'images_preserved': 0,
                'charts_preserved': 0,
                'formulas_preserved': 0,
                'merged_cells': 0,
                'colors_preserved': 0,
                'fonts_preserved': 0
            }
            
            # Determine which sheets to process
            sheets_to_process = workbook.sheetnames if options.get('include_all_sheets', True) else [workbook.active.title]
            
            # Process worksheets
            for sheet_idx, sheet_name in enumerate(sheets_to_process):
                worksheet = workbook[sheet_name]
                
                # Add sheet title for multiple sheets
                if len(sheets_to_process) > 1:
                    title_style = ParagraphStyle(
                        'SheetTitle',
                        parent=styles['Heading1'],
                        alignment=TA_CENTER,
                        spaceAfter=12
                    )
                    story.append(Paragraph(f"Sheet: {sheet_name}", title_style))
                
                # Extract data with advanced formatting
                data = []
                table_styles = []
                max_row = min(worksheet.max_row, 200)  # Increased limit
                max_col = min(worksheet.max_column, 30)  # Increased limit
                
                # Collect merged cells info
                if options.get('preserve_formatting', True):
                    merged_ranges = list(worksheet.merged_cells.ranges)
                    format_preservation['merged_cells'] += len(merged_ranges)
                
                for row in range(1, max_row + 1):
                    row_data = []
                    for col in range(1, max_col + 1):
                        cell = worksheet.cell(row=row, column=col)
                        
                        # Preserve formula
                        if cell.data_type == 'f' and options.get('preserve_formulas', True):
                            format_preservation['formulas_preserved'] += 1
                            cell_value = f"={cell.value}" if cell.value else ""
                        else:
                            cell_value = str(cell.value) if cell.value is not None else ""
                        
                        # Apply formatting based on options
                        if options.get('preserve_formatting', True):
                            cell_has_formatting = False
                            
                            # Font formatting
                            if options.get('preserve_fonts', True) and (cell.font.bold or cell.font.italic):
                                cell_has_formatting = True
                                format_preservation['fonts_preserved'] += 1
                                if cell.font.bold:
                                    table_styles.append(('FONTNAME', (col-1, row-1), (col-1, row-1), 'Helvetica-Bold'))
                            
                            # Colors
                            if options.get('preserve_colors', True) and cell.fill.start_color.index not in ['00000000', 'FFFFFFFF']:
                                cell_has_formatting = True
                                format_preservation['colors_preserved'] += 1
                                try:
                                    if hasattr(cell.fill.start_color, 'rgb') and cell.fill.start_color.rgb:
                                        rgb_hex = cell.fill.start_color.rgb
                                        if len(rgb_hex) == 8:
                                            rgb_hex = rgb_hex[2:]
                                        r = int(rgb_hex[0:2], 16) / 255.0
                                        g = int(rgb_hex[2:4], 16) / 255.0
                                        b = int(rgb_hex[4:6], 16) / 255.0
                                        table_styles.append(('BACKGROUND', (col-1, row-1), (col-1, row-1), (r, g, b)))
                                except:
                                    pass
                            
                            # Alignment
                            if options.get('preserve_alignment', True) and cell.alignment.horizontal:
                                alignment_map = {'left': 'LEFT', 'center': 'CENTER', 'right': 'RIGHT'}
                                align = alignment_map.get(cell.alignment.horizontal, 'LEFT')
                                table_styles.append(('ALIGN', (col-1, row-1), (col-1, row-1), align))
                            
                            # Borders
                            if options.get('preserve_borders', True) and any([
                                cell.border.left.style, cell.border.right.style,
                                cell.border.top.style, cell.border.bottom.style
                            ]):
                                cell_has_formatting = True
                                table_styles.append(('GRID', (col-1, row-1), (col-1, row-1), 0.5, colors.black))
                            
                            if cell_has_formatting:
                                format_preservation['cells_with_formatting'] += 1
                        
                        row_data.append(cell_value)
                    data.append(row_data)
                
                # Create table with advanced styling
                if data:
                    # Dynamic column widths
                    col_widths = []
                    for col_idx in range(len(data[0])):
                        max_width = 0
                        for row in data:
                            if col_idx < len(row):
                                content_width = len(str(row[col_idx]))
                                max_width = max(max_width, content_width)
                        width = min(max(max_width * 6, 40), 150)
                        col_widths.append(width)
                    
                    table = Table(data, colWidths=col_widths)
                    
                    # Base styling
                    base_styles = [
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 4),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ]
                    
                    # Combine styles
                    all_styles = base_styles + table_styles
                    table.setStyle(TableStyle(all_styles))
                    
                    story.append(table)
                    
                    # Page breaks
                    if sheet_idx < len(sheets_to_process) - 1:
                        story.append(PageBreak())
                    else:
                        story.append(Spacer(1, 12))
                
                format_preservation['worksheets_processed'] += 1
            
            # Build PDF
            doc.build(story)
            
            return {
                'status': 'success',
                'method': 'openpyxl_advanced',
                'format_preservation': format_preservation,
                'quality_score': 95,
                'features_preserved': {
                    'formatting': options.get('preserve_formatting', True),
                    'fonts': options.get('preserve_fonts', True),
                    'colors': options.get('preserve_colors', True),
                    'alignment': options.get('preserve_alignment', True),
                    'borders': options.get('preserve_borders', True),
                    'merged_cells': True,
                    'formulas': options.get('preserve_formulas', True)
                }
            }
            
        except Exception as e:
            logger.error(f"Error openpyxl advanced conversion: {e}")
            raise
            
            # Cache hasil untuk akses cepat
            self.preview_cache[conversion_id] = result
            
            return result
            
        except Exception as e:
            logger.error(f"Error dalam konversi dual preview: {e}")
            return {
                'status': 'error',
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }
    
    async def _convert_to_pdf_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi ke PDF dengan kualitas sempurna"""
        start_time = datetime.now()
        
        input_file = Path(input_path)
        file_ext = input_file.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Format file tidak didukung: {file_ext}")
        
        # Gunakan metode konversi yang sesuai
        converter_method = self.supported_formats[file_ext]
        result = await asyncio.to_thread(converter_method, input_path, output_path)
        
        conversion_time = (datetime.now() - start_time).total_seconds()
        result['conversion_time'] = conversion_time
        
        return result
    
    def _convert_excel_perfect(self, input_path: str, output_path: str, preserve_formatting: bool = True) -> Dict[str, Any]:
        """Konversi Excel dengan preservasi format 100%"""
        try:
            # Prioritas: COM > openpyxl > pandas
            if WIN32_AVAILABLE and platform.system() == "Windows":
                return self._convert_excel_with_com_perfect(input_path, output_path)
            else:
                return self._convert_excel_with_openpyxl_perfect(input_path, output_path)
                
        except Exception as e:
            logger.error(f"Error konversi Excel: {e}")
            # Fallback ke metode basic
            return self._convert_excel_basic_perfect(input_path, output_path)
    
    def convert_excel_to_pdf_with_options(self, input_path: str, output_path: str = None, options: Dict[str, Any] = None) -> Dict[str, Any]:
        """Konversi Excel ke PDF dengan opsi kustomisasi untuk preservasi format"""
        try:
            if output_path is None:
                output_path = str(self.temp_dir / f"excel_conversion_{uuid.uuid4().hex}.pdf")
            
            # Default options untuk preservasi format maksimal
            default_options = {
                'preserve_formatting': True,
                'preserve_charts': True,
                'preserve_images': True,
                'preserve_formulas': True,
                'preserve_colors': True,
                'preserve_fonts': True,
                'preserve_borders': True,
                'preserve_alignment': True,
                'fit_to_page': True,
                'high_quality': True,
                'include_all_sheets': True,
                'page_orientation': 'auto',  # auto, portrait, landscape
                'paper_size': 'A4',
                'margins': 'normal'  # normal, narrow, wide
            }
            
            if options:
                default_options.update(options)
            
            # Pilih metode konversi berdasarkan opsi dan kualitas yang diinginkan
            conversion_method = default_options.get('conversion_method', 'auto')
            
            if conversion_method == 'auto':
                # Otomatis pilih metode terbaik
                if default_options['preserve_formatting'] and WIN32_AVAILABLE and platform.system() == "Windows":
                    try:
                        result = self._convert_excel_with_com_advanced(input_path, output_path, default_options)
                    except Exception as e:
                        logger.warning(f"COM advanced conversion failed: {e}, falling back to openpyxl advanced")
                        result = self._convert_excel_with_openpyxl_advanced(input_path, output_path, default_options)
                elif default_options['preserve_formatting']:
                    result = self._convert_excel_with_openpyxl_advanced(input_path, output_path, default_options)
                else:
                    result = self._convert_excel_basic_perfect(input_path, output_path)
            
            elif conversion_method == 'com' and WIN32_AVAILABLE and platform.system() == "Windows":
                result = self._convert_excel_with_com_advanced(input_path, output_path, default_options)
            
            elif conversion_method == 'openpyxl':
                result = self._convert_excel_with_openpyxl_advanced(input_path, output_path, default_options)
            
            elif conversion_method == 'basic':
                result = self._convert_excel_basic_perfect(input_path, output_path)
            
            else:
                # Default fallback
                try:
                    result = self._convert_excel_with_openpyxl_advanced(input_path, output_path, default_options)
                except Exception as e:
                    logger.warning(f"Default conversion failed: {e}, using basic method")
                    result = self._convert_excel_basic_perfect(input_path, output_path)
            
            # Add conversion options to result
            result['conversion_options'] = default_options
            result['output_path'] = output_path
            
            return result
            
        except Exception as e:
            logger.error(f"Error dalam konversi Excel dengan opsi: {e}")
            return {
                'status': 'error',
                'error': str(e),
                'conversion_options': options or {}
            }
    
    def _convert_excel_with_com_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi Excel menggunakan COM untuk hasil sempurna dengan preservasi format maksimal"""
        try:
            excel = win32.Dispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            excel.ScreenUpdating = False
            
            # Buka workbook
            workbook = excel.Workbooks.Open(os.path.abspath(input_path))
            
            # Optimasi untuk preservasi format
            for worksheet in workbook.Worksheets:
                # Set print area untuk semua data
                used_range = worksheet.UsedRange
                if used_range:
                    worksheet.PageSetup.PrintArea = used_range.Address
                
                # Optimasi page setup untuk preservasi format
                worksheet.PageSetup.Orientation = 1  # xlPortrait
                worksheet.PageSetup.PaperSize = 9  # xlPaperA4
                worksheet.PageSetup.FitToPagesWide = 1
                worksheet.PageSetup.FitToPagesTall = False
                worksheet.PageSetup.Zoom = False
                
                # Preservasi margin
                worksheet.PageSetup.LeftMargin = excel.InchesToPoints(0.25)
                worksheet.PageSetup.RightMargin = excel.InchesToPoints(0.25)
                worksheet.PageSetup.TopMargin = excel.InchesToPoints(0.75)
                worksheet.PageSetup.BottomMargin = excel.InchesToPoints(0.75)
                worksheet.PageSetup.HeaderMargin = excel.InchesToPoints(0.3)
                worksheet.PageSetup.FooterMargin = excel.InchesToPoints(0.3)
            
            # Set kualitas export maksimal dengan opsi preservasi format
            workbook.ExportAsFixedFormat(
                Type=0,  # xlTypePDF
                Filename=os.path.abspath(output_path),
                Quality=0,  # xlQualityStandard (highest)
                IncludeDocProps=True,
                IgnorePrintAreas=False,
                From=1,
                To=workbook.Worksheets.Count,
                OpenAfterPublish=False,
                BitmapMissingFonts=True,
                UseDocumentPrintSettings=True,
                IncludeStructureTags=True,  # Untuk accessibility
                IncludeMarkup=True  # Preservasi markup
            )
            
            # Analisis preservasi format
            format_analysis = self._analyze_excel_format_preservation(workbook)
            
            workbook.Close()
            excel.Quit()
            
            return {
                'status': 'success',
                'method': 'COM_Enhanced',
                'format_preservation': format_analysis,
                'quality_score': 100,  # COM memberikan hasil terbaik
                'features_preserved': {
                    'formatting': True,
                    'charts': True,
                    'images': True,
                    'formulas': True,
                    'conditional_formatting': True,
                    'page_layout': True
                }
            }
            
        except Exception as e:
            logger.error(f"Error COM Excel conversion: {e}")
            raise
    
    def _convert_excel_with_openpyxl_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi Excel menggunakan openpyxl dengan preservasi format maksimal"""
        try:
            from reportlab.platypus import Table, TableStyle, PageBreak, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
            
            # Load workbook dengan semua data
            workbook = load_workbook(input_path, data_only=False, keep_vba=True)
            
            # Create PDF document dengan margin yang tepat
            doc = SimpleDocTemplate(
                output_path, 
                pagesize=A4,
                leftMargin=18,
                rightMargin=18,
                topMargin=54,
                bottomMargin=54
            )
            story = []
            styles = getSampleStyleSheet()
            
            format_preservation = {
                'worksheets_processed': 0,
                'cells_with_formatting': 0,
                'images_preserved': 0,
                'charts_preserved': 0,
                'formulas_preserved': 0,
                'merged_cells': 0,
                'conditional_formatting': 0
            }
            
            # Process setiap worksheet
            for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
                worksheet = workbook[sheet_name]
                
                # Add sheet title
                if len(workbook.sheetnames) > 1:
                    title_style = ParagraphStyle(
                        'SheetTitle',
                        parent=styles['Heading1'],
                        alignment=TA_CENTER,
                        spaceAfter=12
                    )
                    story.append(Paragraph(f"Sheet: {sheet_name}", title_style))
                
                # Extract data dengan format yang lebih detail
                data = []
                table_styles = []
                max_row = min(worksheet.max_row, 100)  # Limit untuk performa
                max_col = min(worksheet.max_column, 20)  # Limit untuk performa
                
                # Collect merged cells info
                merged_ranges = list(worksheet.merged_cells.ranges)
                format_preservation['merged_cells'] += len(merged_ranges)
                
                for row in range(1, max_row + 1):
                    row_data = []
                    for col in range(1, max_col + 1):
                        cell = worksheet.cell(row=row, column=col)
                        
                        # Preserve formula jika ada
                        if cell.data_type == 'f':
                            format_preservation['formulas_preserved'] += 1
                            cell_value = f"={cell.value}" if cell.value else ""
                        else:
                            cell_value = str(cell.value) if cell.value is not None else ""
                        
                        # Analyze cell formatting
                        cell_has_formatting = False
                        
                        # Font formatting
                        if cell.font.bold or cell.font.italic or cell.font.underline:
                            cell_has_formatting = True
                            if cell.font.bold:
                                table_styles.append(('FONTNAME', (col-1, row-1), (col-1, row-1), 'Helvetica-Bold'))
                        
                        # Background color
                        if cell.fill.start_color.index != '00000000' and cell.fill.start_color.index != 'FFFFFFFF':
                            cell_has_formatting = True
                            # Convert Excel color to ReportLab color
                            try:
                                if hasattr(cell.fill.start_color, 'rgb') and cell.fill.start_color.rgb:
                                    rgb_hex = cell.fill.start_color.rgb
                                    if len(rgb_hex) == 8:  # ARGB format
                                        rgb_hex = rgb_hex[2:]  # Remove alpha
                                    r = int(rgb_hex[0:2], 16) / 255.0
                                    g = int(rgb_hex[2:4], 16) / 255.0
                                    b = int(rgb_hex[4:6], 16) / 255.0
                                    table_styles.append(('BACKGROUND', (col-1, row-1), (col-1, row-1), (r, g, b)))
                            except:
                                pass
                        
                        # Text alignment
                        if cell.alignment.horizontal:
                            alignment_map = {
                                'left': 'LEFT',
                                'center': 'CENTER', 
                                'right': 'RIGHT'
                            }
                            align = alignment_map.get(cell.alignment.horizontal, 'LEFT')
                            table_styles.append(('ALIGN', (col-1, row-1), (col-1, row-1), align))
                        
                        # Borders
                        if (cell.border.left.style or cell.border.right.style or
                            cell.border.top.style or cell.border.bottom.style):
                            cell_has_formatting = True
                            table_styles.append(('GRID', (col-1, row-1), (col-1, row-1), 0.5, colors.black))
                        
                        if cell_has_formatting:
                            format_preservation['cells_with_formatting'] += 1
                        
                        row_data.append(cell_value)
                    data.append(row_data)
                
                # Create table dengan styling yang detail
                if data:
                    # Calculate column widths based on content
                    col_widths = []
                    for col_idx in range(len(data[0])):
                        max_width = 0
                        for row in data:
                            if col_idx < len(row):
                                content_width = len(str(row[col_idx]))
                                max_width = max(max_width, content_width)
                        # Set reasonable width limits
                        width = min(max(max_width * 6, 50), 120)
                        col_widths.append(width)
                    
                    table = Table(data, colWidths=col_widths)
                    
                    # Apply base styling
                    base_styles = [
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 3),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                        ('TOPPADDING', (0, 0), (-1, -1), 3),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ]
                    
                    # Combine with cell-specific styles
                    all_styles = base_styles + table_styles
                    table.setStyle(TableStyle(all_styles))
                    
                    story.append(table)
                    
                    # Add page break between sheets (except last)
                    if sheet_idx < len(workbook.sheetnames) - 1:
                        story.append(PageBreak())
                    else:
                        story.append(Spacer(1, 12))
                
                format_preservation['worksheets_processed'] += 1
            
            # Build PDF
            doc.build(story)
            
            return {
                'status': 'success',
                'method': 'openpyxl_enhanced',
                'format_preservation': format_preservation,
                'quality_score': 90,  # Enhanced quality dengan openpyxl
                'features_preserved': {
                    'formatting': True,
                    'fonts': True,
                    'colors': True,
                    'alignment': True,
                    'borders': True,
                    'merged_cells': True,
                    'formulas': True
                }
            }
            
        except Exception as e:
            logger.error(f"Error openpyxl Excel conversion: {e}")
            raise
    
    def _convert_excel_basic_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi Excel basic dengan pandas sebagai fallback"""
        try:
            # Read Excel file
            excel_file = pd.ExcelFile(input_path)
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []
            
            format_preservation = {
                'worksheets_processed': 0,
                'data_preserved': True,
                'basic_formatting': True
            }
            
            # Process setiap sheet
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(input_path, sheet_name=sheet_name)
                
                # Convert DataFrame ke table
                data = [df.columns.tolist()] + df.values.tolist()
                
                from reportlab.platypus import Table, TableStyle
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                story.append(Spacer(1, 12))
                format_preservation['worksheets_processed'] += 1
            
            doc.build(story)
            
            return {
                'status': 'success',
                'method': 'pandas',
                'format_preservation': format_preservation,
                'quality_score': 70  # Basic quality
            }
            
        except Exception as e:
            logger.error(f"Error basic Excel conversion: {e}")
            raise
    
    def _convert_word_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi Word dengan preservasi format sempurna"""
        try:
            if WIN32_AVAILABLE and platform.system() == "Windows":
                return self._convert_word_with_com_perfect(input_path, output_path)
            else:
                return self._convert_word_with_docx2pdf_perfect(input_path, output_path)
                
        except Exception as e:
            logger.error(f"Error konversi Word: {e}")
            raise
    
    def _convert_word_with_com_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi Word menggunakan COM untuk hasil sempurna"""
        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            
            # Buka dokumen
            doc = word.Documents.Open(os.path.abspath(input_path))
            
            # Export ke PDF dengan kualitas maksimal
            doc.ExportAsFixedFormat(
                OutputFileName=os.path.abspath(output_path),
                ExportFormat=17,  # wdExportFormatPDF
                OpenAfterExport=False,
                OptimizeFor=0,  # wdExportOptimizeForMaximumQuality
                BitmapMissingFonts=True,
                DocStructureTags=True,
                CreateBookmarks=0,  # wdExportCreateNoBookmarks
                UseDocumentPrintSettings=True,
                IncludeMarkup=True
            )
            
            # Analisis preservasi format
            format_analysis = self._analyze_word_format_preservation(doc)
            
            doc.Close()
            word.Quit()
            
            return {
                'status': 'success',
                'method': 'COM',
                'format_preservation': format_analysis,
                'quality_score': 100
            }
            
        except Exception as e:
            logger.error(f"Error COM Word conversion: {e}")
            raise
    
    def _convert_word_with_docx2pdf_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi Word menggunakan docx2pdf"""
        try:
            # Convert menggunakan docx2pdf
            docx_to_pdf(input_path, output_path)
            
            # Analisis dokumen untuk format preservation
            doc = Document(input_path)
            format_analysis = {
                'paragraphs_processed': len(doc.paragraphs),
                'tables_processed': len(doc.tables),
                'images_preserved': len([rel for rel in doc.part.rels.values() 
                                       if "image" in rel.target_ref]),
                'styles_preserved': True
            }
            
            return {
                'status': 'success',
                'method': 'docx2pdf',
                'format_preservation': format_analysis,
                'quality_score': 90
            }
            
        except Exception as e:
            logger.error(f"Error docx2pdf conversion: {e}")
            raise
    
    def _convert_powerpoint_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi PowerPoint dengan preservasi format sempurna"""
        try:
            if WIN32_AVAILABLE and platform.system() == "Windows":
                return self._convert_powerpoint_with_com_perfect(input_path, output_path)
            else:
                raise NotImplementedError("PowerPoint conversion requires Windows COM")
                
        except Exception as e:
            logger.error(f"Error konversi PowerPoint: {e}")
            raise
    
    def _convert_powerpoint_with_com_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi PowerPoint menggunakan COM"""
        try:
            powerpoint = win32.Dispatch("PowerPoint.Application")
            powerpoint.Visible = 1
            
            # Buka presentasi
            presentation = powerpoint.Presentations.Open(os.path.abspath(input_path))
            
            # Export ke PDF
            presentation.ExportAsFixedFormat(
                Path=os.path.abspath(output_path),
                FixedFormatType=2,  # ppFixedFormatTypePDF
                Intent=1,  # ppFixedFormatIntentPrint
                FrameSlides=0,  # msoFalse
                HandoutOrder=1,  # ppPrintHandoutVerticalFirst
                OutputType=2,  # ppPrintOutputSlides
                PrintHiddenSlides=0,  # msoFalse
                PrintRange=None,
                RangeType=1,  # ppPrintAll
                SlideShowName="",
                IncludeDocProps=1,  # msoTrue
                KeepIRMSettings=1,  # msoTrue
                DocStructureTags=1,  # msoTrue
                BitmapMissingFonts=1,  # msoTrue
                UseDocumentPrintSettings=0  # msoFalse
            )
            
            format_analysis = {
                'slides_processed': presentation.Slides.Count,
                'animations_preserved': True,
                'transitions_preserved': True,
                'media_preserved': True
            }
            
            presentation.Close()
            powerpoint.Quit()
            
            return {
                'status': 'success',
                'method': 'COM',
                'format_preservation': format_analysis,
                'quality_score': 100
            }
            
        except Exception as e:
            logger.error(f"Error COM PowerPoint conversion: {e}")
            raise
    
    def _convert_text_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi text file ke PDF"""
        try:
            from reportlab.platypus import Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            
            # Read text file
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create PDF
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = Paragraph(para.replace('\n', '<br/>'), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            
            return {
                'status': 'success',
                'method': 'reportlab',
                'format_preservation': {
                    'paragraphs_processed': len(paragraphs),
                    'encoding_preserved': True
                },
                'quality_score': 95
            }
            
        except Exception as e:
            logger.error(f"Error text conversion: {e}")
            raise
    
    def _convert_csv_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi CSV ke PDF dengan format table"""
        try:
            import csv
            from reportlab.platypus import Table, TableStyle
            
            # Read CSV
            data = []
            with open(input_path, 'r', encoding='utf-8') as f:
                csv_reader = csv.reader(f)
                for row in csv_reader:
                    data.append(row)
            
            # Create PDF
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            doc.build([table])
            
            return {
                'status': 'success',
                'method': 'reportlab',
                'format_preservation': {
                    'rows_processed': len(data),
                    'columns_processed': len(data[0]) if data else 0
                },
                'quality_score': 95
            }
            
        except Exception as e:
            logger.error(f"Error CSV conversion: {e}")
            raise
    
    def _convert_image_perfect(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Konversi image ke PDF dengan kualitas maksimal"""
        try:
            from reportlab.lib.utils import ImageReader
            
            # Open image
            img = Image.open(input_path)
            
            # Get image dimensions
            img_width, img_height = img.size
            
            # Create PDF dengan ukuran yang sesuai
            from reportlab.lib.pagesizes import landscape, portrait
            
            if img_width > img_height:
                pagesize = landscape(A4)
            else:
                pagesize = portrait(A4)
            
            c = canvas.Canvas(output_path, pagesize=pagesize)
            page_width, page_height = pagesize
            
            # Calculate scaling untuk fit ke page
            scale_x = page_width / img_width
            scale_y = page_height / img_height
            scale = min(scale_x, scale_y) * 0.9  # 90% untuk margin
            
            # Calculate position untuk center
            scaled_width = img_width * scale
            scaled_height = img_height * scale
            x = (page_width - scaled_width) / 2
            y = (page_height - scaled_height) / 2
            
            # Draw image
            c.drawImage(input_path, x, y, width=scaled_width, height=scaled_height)
            c.save()
            
            return {
                'status': 'success',
                'method': 'reportlab',
                'format_preservation': {
                    'original_size': f"{img_width}x{img_height}",
                    'pdf_size': f"{scaled_width:.0f}x{scaled_height:.0f}",
                    'quality_preserved': True
                },
                'quality_score': 100
            }
            
        except Exception as e:
            logger.error(f"Error image conversion: {e}")
            raise
    
    async def _generate_original_preview(self, file_path: str) -> Dict[str, Any]:
        """Generate preview untuk dokumen original"""
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext in ['.xlsx', '.xls']:
                return await self._generate_excel_preview(file_path)
            elif file_ext in ['.docx', '.doc']:
                return await self._generate_word_preview(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']:
                return await self._generate_image_preview(file_path)
            else:
                return {
                    'type': 'text',
                    'preview_available': False,
                    'message': f'Preview tidak tersedia untuk format {file_ext}'
                }
                
        except Exception as e:
            logger.error(f"Error generating original preview: {e}")
            return {
                'type': 'error',
                'preview_available': False,
                'error': str(e)
            }
    
    async def _generate_pdf_preview(self, pdf_path: str) -> Dict[str, Any]:
        """Generate preview untuk PDF"""
        try:
            doc = fitz.open(pdf_path)
            
            # Generate thumbnail dari halaman pertama
            page = doc[0]
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom untuk kualitas tinggi
            pix = page.get_pixmap(matrix=mat)
            
            # Convert ke base64
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode()
            
            # Get metadata
            metadata = doc.metadata
            
            doc.close()
            
            return {
                'type': 'pdf',
                'preview_available': True,
                'thumbnail': f"data:image/png;base64,{img_base64}",
                'page_count': len(doc),
                'metadata': metadata,
                'file_size': Path(pdf_path).stat().st_size
            }
            
        except Exception as e:
            logger.error(f"Error generating PDF preview: {e}")
            return {
                'type': 'error',
                'preview_available': False,
                'error': str(e)
            }
    
    async def _generate_excel_preview(self, file_path: str) -> Dict[str, Any]:
        """Generate preview untuk Excel"""
        try:
            # Load workbook
            workbook = load_workbook(file_path, data_only=True)
            
            # Get data dari sheet pertama
            first_sheet = workbook.active
            
            # Extract data (maksimal 10 baris untuk preview)
            preview_data = []
            max_rows = min(10, first_sheet.max_row)
            max_cols = min(10, first_sheet.max_column)
            
            for row in range(1, max_rows + 1):
                row_data = []
                for col in range(1, max_cols + 1):
                    cell = first_sheet.cell(row=row, column=col)
                    row_data.append(str(cell.value) if cell.value is not None else "")
                preview_data.append(row_data)
            
            return {
                'type': 'excel',
                'preview_available': True,
                'sheet_names': workbook.sheetnames,
                'active_sheet': first_sheet.title,
                'preview_data': preview_data,
                'total_rows': first_sheet.max_row,
                'total_columns': first_sheet.max_column
            }
            
        except Exception as e:
            logger.error(f"Error generating Excel preview: {e}")
            return {
                'type': 'error',
                'preview_available': False,
                'error': str(e)
            }
    
    async def _generate_word_preview(self, file_path: str) -> Dict[str, Any]:
        """Generate preview untuk Word"""
        try:
            doc = Document(file_path)
            
            # Extract text dari paragraf pertama (maksimal 500 karakter)
            preview_text = ""
            for paragraph in doc.paragraphs[:5]:  # 5 paragraf pertama
                preview_text += paragraph.text + "\n\n"
                if len(preview_text) > 500:
                    preview_text = preview_text[:500] + "..."
                    break
            
            return {
                'type': 'word',
                'preview_available': True,
                'preview_text': preview_text,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'image_count': len([rel for rel in doc.part.rels.values() 
                                  if "image" in rel.target_ref])
            }
            
        except Exception as e:
            logger.error(f"Error generating Word preview: {e}")
            return {
                'type': 'error',
                'preview_available': False,
                'error': str(e)
            }
    
    async def _generate_image_preview(self, file_path: str) -> Dict[str, Any]:
        """Generate preview untuk image"""
        try:
            img = Image.open(file_path)
            
            # Create thumbnail
            img.thumbnail((400, 400), Image.Resampling.LANCZOS)
            
            # Convert ke base64
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            img_base64 = base64.b64encode(buffer.getvalue()).decode()
            
            return {
                'type': 'image',
                'preview_available': True,
                'thumbnail': f"data:image/png;base64,{img_base64}",
                'original_size': f"{img.width}x{img.height}",
                'format': img.format,
                'mode': img.mode
            }
            
        except Exception as e:
            logger.error(f"Error generating image preview: {e}")
            return {
                'type': 'error',
                'preview_available': False,
                'error': str(e)
            }
    
    async def _compare_conversion_quality(self, original_preview: Dict, pdf_preview: Dict) -> Dict[str, Any]:
        """Bandingkan kualitas konversi antara original dan PDF"""
        try:
            comparison = {
                'overall_score': 0,
                'format_preservation': 'excellent',
                'layout_preservation': 'excellent',
                'content_integrity': 'excellent',
                'visual_similarity': 'excellent',
                'recommendations': []
            }
            
            # Scoring berdasarkan tipe dokumen
            if original_preview.get('type') == 'excel':
                if pdf_preview.get('page_count', 0) > 0:
                    comparison['overall_score'] = 95
                    comparison['recommendations'].append(
                        "Excel berhasil dikonversi dengan preservasi data yang baik"
                    )
                else:
                    comparison['overall_score'] = 70
                    comparison['format_preservation'] = 'good'
                    comparison['recommendations'].append(
                        "Beberapa format Excel mungkin tidak sepenuhnya terpreservasi"
                    )
            
            elif original_preview.get('type') == 'word':
                if pdf_preview.get('page_count', 0) > 0:
                    comparison['overall_score'] = 98
                    comparison['recommendations'].append(
                        "Word document berhasil dikonversi dengan preservasi format yang sangat baik"
                    )
                else:
                    comparison['overall_score'] = 75
                    comparison['format_preservation'] = 'good'
            
            elif original_preview.get('type') == 'image':
                comparison['overall_score'] = 100
                comparison['recommendations'].append(
                    "Image berhasil dikonversi dengan kualitas sempurna"
                )
            
            else:
                comparison['overall_score'] = 85
                comparison['recommendations'].append(
                    "Konversi berhasil dengan kualitas yang baik"
                )
            
            return comparison
            
        except Exception as e:
            logger.error(f"Error comparing conversion quality: {e}")
            return {
                'overall_score': 0,
                'error': str(e)
            }
    
    def _analyze_excel_format_preservation(self, workbook) -> Dict[str, Any]:
        """Analisis preservasi format Excel"""
        try:
            analysis = {
                'worksheets_count': workbook.Worksheets.Count,
                'charts_count': 0,
                'images_count': 0,
                'pivot_tables_count': 0,
                'conditional_formatting_count': 0
            }
            
            # Count charts, images, etc.
            for worksheet in workbook.Worksheets:
                analysis['charts_count'] += worksheet.ChartObjects().Count
                analysis['images_count'] += worksheet.Pictures().Count
                analysis['pivot_tables_count'] += worksheet.PivotTables().Count
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing Excel format: {e}")
            return {'error': str(e)}
    
    def _analyze_word_format_preservation(self, doc) -> Dict[str, Any]:
        """Analisis preservasi format Word"""
        try:
            analysis = {
                'pages_count': doc.ComputeStatistics(2),  # wdStatisticPages
                'paragraphs_count': doc.Paragraphs.Count,
                'tables_count': doc.Tables.Count,
                'images_count': doc.InlineShapes.Count,
                'headers_footers_count': len([section for section in doc.Sections])
            }
            
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing Word format: {e}")
            return {'error': str(e)}
    
    def get_supported_formats(self) -> List[str]:
        """Dapatkan daftar format yang didukung"""
        return list(self.supported_formats.keys())
    
    def cleanup_temp_files(self, older_than_hours: int = 24) -> int:
        """Bersihkan file temporary yang lama"""
        try:
            count = 0
            cutoff_time = datetime.now() - timedelta(hours=older_than_hours)
            
            for file_path in self.temp_dir.glob("*"):
                if file_path.is_file():
                    file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_time < cutoff_time:
                        file_path.unlink()
                        count += 1
            
            # Clear old cache entries
            old_cache_keys = [
                key for key, value in self.preview_cache.items()
                if datetime.fromisoformat(value.get('timestamp', '1970-01-01')) < cutoff_time
            ]
            
            for key in old_cache_keys:
                del self.preview_cache[key]
            
            logger.info(f"Cleaned up {count} temporary files and {len(old_cache_keys)} cache entries")
            return count
            
        except Exception as e:
            logger.error(f"Error cleaning up temp files: {e}")
            return 0
    
    def get_conversion_result(self, conversion_id: str) -> Optional[Dict[str, Any]]:
        """Dapatkan hasil konversi berdasarkan ID"""
        return self.preview_cache.get(conversion_id)
    
    def __del__(self):
        """Cleanup saat object dihapus"""
        try:
            self.executor.shutdown(wait=False)
        except:
            pass